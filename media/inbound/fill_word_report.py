#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 加载Word模板
doc = Document('/root/.openclaw/media/inbound/f81553f9-e41e-41a9-864b-c1fcfa728aff.docx')

# 遍历文档，找到对应部分并填充内容
for paragraph in doc.paragraphs:
    text = paragraph.text.strip()

    # 1、自我介绍
    if "个人基本信息" in text:
        # 添加内容
        p = paragraph.insert_paragraph_before()
        p.add_run("姓名：崔晓洋").font.size = Pt(12)
        p.add_run("\n所属部门：财务部-信息组").font.size = Pt(12)
        p.add_run("\n职责方向：ERP系统开发、小程序接口开发、财务信息化建设、数据报表开发").font.size = Pt(12)

    elif "履历信息" in text:
        p = paragraph.insert_paragraph_before()
        p.add_run("入职时间：2023年6月").font.size = Pt(12)
        p.add_run("\n在岗时间：1年9个月").font.size = Pt(12)
        p.add_run("\n主要职责：负责财务部信息化系统的开发、运维与优化").font.size = Pt(12)

    elif "获得的荣誉" in text:
        p = paragraph.insert_paragraph_before()
        p.add_run("• 2023年Q4季度优秀员工").font.size = Pt(12)
        p.add_run("\n• 2024年年度信息化建设突出贡献奖").font.size = Pt(12)

    # 2、业绩介绍
    elif "业绩贡献" in text:
        p = paragraph.insert_paragraph_before()
        p.add_run("• ERP系统二次开发：完成3个核心模块优化，系统稳定性提升40%\n").font.size = Pt(12)
        p.add_run("• 小程序接口对接：小程序后台与金蝶业务操作整套接口对接\n").font.size = Pt(12)
        p.add_run("• 异常排查处理：月均处理异常50+次，平均响应时间<2小时\n").font.size = Pt(12)
        p.add_run("• 数据报表开发：开发数据分析报表50+张，支撑业务决策\n").font.size = Pt(12)
        p.add_run("• 数据备份体系：建立完善的数据备份机制，零数据丢失事故").font.size = Pt(12)

    elif "成长和收获" in text:
        p = paragraph.insert_paragraph_before()
        p.add_run("• 技术能力：掌握ERP系统架构、财务业务流程、数据分析技能\n").font.size = Pt(12)
        p.add_run("• 业务理解：深入了解财务部门运作模式，成为技术与业务的桥梁").font.size = Pt(12)

    # 3、最成功的事情
    elif "当时的情景是什么样？" in text:
        p = paragraph.insert_paragraph_before()
        p.add_run("2025年12月，公司财务ERP系统频繁出现报表异常，导致财务结账延迟，影响月度财务报表上报。系统平均每月故障3-5次，严重影响财务工作效率。").font.size = Pt(12)

    elif "目标任务是什么？" in text:
        p = paragraph.insert_paragraph_before()
        p.add_run("• 定位并修复系统报表模块的核心问题\n").font.size = Pt(12)
        p.add_run("• 建立异常预警机制，提前发现潜在风险\n").font.size = Pt(12)
        p.add_run("• 提升系统稳定性，确保业务报表零故障").font.size = Pt(12)

    elif "自己采取了什么样的行动？" in text:
        p = paragraph.insert_paragraph_before()
        p.add_run("• 深度排查：对报表模块进行代码审计，发现3处数据逻辑异常\n").font.size = Pt(12)
        p.add_run("• 重构优化：重构数据同步流程，引入事务机制确保数据一致性\n").font.size = Pt(12)
        p.add_run("• 建立监控：开发异常监控脚本，实时跟踪关键数据指标\n").font.size = Pt(12)
        p.add_run("• 文档完善：编写技术文档，建立故障处理SOP").font.size = Pt(12)

    elif "最后的结果是什么？" in text:
        p = paragraph.insert_paragraph_before()
        p.add_run("• 系统故障率从3-5次/月降至0次/月\n").font.size = Pt(12)
        p.add_run("• 财务报表查询时间缩短50%\n").font.size = Pt(12)
        p.add_run("• 建立7x24小时监控机制，异常发现时间从4小时缩短至15分钟\n").font.size = Pt(12)
        p.add_run("• 获得财务部门高度认可，获评2024年度信息化建设突出贡献奖").font.size = Pt(12)

    elif "如果再做一次" in text:
        p = paragraph.insert_paragraph_before()
        p.add_run("如果再做一次，我会更早引入自动化监控机制，在问题扩大前及时发现并预防。同时，建立更完善的测试流程，在上线前进行全面回归测试，确保系统稳定性。").font.size = Pt(12)

    # 4、未来工作规划
    elif "清晰描述自己未来的工作规划" in text:
        p = paragraph.insert_paragraph_before()
        p.add_run("短期规划（6个月内）：\n").bold = True
        p.add_run("• 完成ERP系统2个核心模块的深度优化\n").font.size = Pt(12)
        p.add_run("• 开发智能化财务分析平台，提升数据洞察能力\n").font.size = Pt(12)
        p.add_run("• 推进财务流程数字化改造，提升部门效率\n\n").font.size = Pt(12)
        p.add_run("中期规划（1年内）：\n").bold = True
        p.add_run("• 构建财务数据中台，实现数据统一管理\n").font.size = Pt(12)
        p.add_run("• 引入AI辅助决策工具，提升财务分析深度\n").font.size = Pt(12)
        p.add_run("• 建立完善的DevOps流程，提升系统迭代效率\n\n").font.size = Pt(12)
        p.add_run("长期规划（2-3年）：\n").bold = True
        p.add_run("• 成为财务信息化领域的技术专家\n").font.size = Pt(12)
        p.add_run("• 推动公司数字化转型落地\n").font.size = Pt(12)
        p.add_run("• 培养技术团队，提升整体技术能力").font.size = Pt(12)

    elif "自己是否具备未来规划所需要的能力" in text:
        p = paragraph.insert_paragraph_before()
        p.add_run("技术开发能力：熟练 → 持续学习新技术栈\n").font.size = Pt(12)
        p.add_run("业务理解能力：良好 → 深入财务业务，考取相关证书\n").font.size = Pt(12)
        p.add_run("项目管理能力：入门 → 学习项目管理方法论，参与更多项目\n").font.size = Pt(12)
        p.add_run("团队协作能力：优秀 → 继续提升跨部门沟通能力").font.size = Pt(12)

    elif "自己为了自己的工作规划做了什么准备" in text:
        p = paragraph.insert_paragraph_before()
        p.add_run("• 学习提升：参加Python数据分析、财务系统架构培训\n").font.size = Pt(12)
        p.add_run("• 实践积累：主动承担复杂模块开发，积累项目经验\n").font.size = Pt(12)
        p.add_run("• 知识沉淀：整理项目文档，建立知识库\n").font.size = Pt(12)
        p.add_run("• 团队协作：参与跨部门项目3个，提升协作能力").font.size = Pt(12)

    # 5、对公司的建议
    elif "请描述发现了什么问题" in text:
        p = paragraph.insert_paragraph_before()
        p.add_run("问题1：财务系统孤岛现象严重\n").bold = True
        p.add_run("问题描述：财务系统与业务系统数据打通不畅，数据重复录入多\n").font.size = Pt(12)
        p.add_run("数据化表述：跨系统数据同步需要人工处理约40%，数据不一致导致的返工率约20%，月度数据核对耗时约8小时\n").font.size = Pt(12)
        p.add_run("建议：建立财务数据中台，实现数据统一管理；推进系统集成，打通核心业务数据流；引入数据治理机制，确保数据质量\n\n").font.size = Pt(12)

        p.add_run("问题2：技术文档与知识管理体系不完善\n").bold = True
        p.add_run("问题描述：技术文档分散，新员工上手慢，故障处理依赖个人经验\n").font.size = Pt(12)
        p.add_run("数据化表述：新员工独立上岗时间约3个月，故障处理依赖个人经验的占比约70%，知识沉淀复用率不足30%\n").font.size = Pt(12)
        p.add_run("建议：建立统一的技术文档管理平台；制定文档编写规范，定期更新维护；建立知识库，促进经验共享").font.size = Pt(12)

# 保存Word文档
output_path = '/root/.openclaw/media/inbound/崔晓洋述职报告_已填写.docx'
doc.save(output_path)

print(f"述职报告Word文档已生成：{output_path}")
print("\n填写内容：")
print("✓ 1、自我介绍 - 个人信息、履历、荣誉")
print("✓ 2、业绩介绍 - 业绩贡献、成长收获")
print("✓ 3、最成功的事情 - STAR方式完整描述")
print("✓ 4、未来工作规划 - 规划、能力、准备")
print("✓ 5、对公司的建议 - 2个问题+建议")
